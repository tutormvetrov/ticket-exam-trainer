from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

from application.facade import AppFacade
from application.import_service import DocumentImportService, TicketCandidate
from application.settings import DEFAULT_OLLAMA_SETTINGS
from application.settings_store import SettingsStore
from domain.answer_profile import AnswerBlockCode, AnswerProfileCode
from domain.knowledge import Exam, Section
from infrastructure.db import connect_initialized, get_database_path


SOURCE_TEXT = """Section 1. Public assets

Ticket 1. What is public property as an object of management? Public property is a public resource assigned to public bodies. Examples include land, buildings and infrastructure. The asset has a legal regime and requires control. The management cycle includes accounting, valuation, use and review.

Ticket 2. How is efficiency of public property evaluated? Efficiency is evaluated through public goals, usage results and cost control. For example, analysts check utilization and social effect.
"""


def _build_facade(tmp_path: Path, *, import_llm_assist: bool = False) -> AppFacade:
    workspace_root = tmp_path / "workspace"
    workspace_root.mkdir(parents=True, exist_ok=True)
    database_path = get_database_path(workspace_root)
    connection = connect_initialized(database_path)
    settings_store = SettingsStore(workspace_root / "app_data" / "settings.json")
    settings_store.save(
        replace(
            DEFAULT_OLLAMA_SETTINGS,
            auto_check_ollama_on_start=False,
            auto_check_updates_on_start=False,
            import_llm_assist=import_llm_assist,
            examiner_followups=False,
            rewrite_questions=False,
        )
    )
    return AppFacade(workspace_root, connection, settings_store)


def test_build_ticket_model_from_text() -> None:
    service = DocumentImportService()
    candidate = TicketCandidate(
        index=1,
        title="What is public property as an object of management?",
        body="Public property is a public resource. Examples include land and buildings. The management cycle includes accounting and review.",
        confidence=0.9,
        section_title="public-assets",
    )
    ticket, used_llm, warning = service.build_ticket_map(candidate, "exam-demo", "public-assets", "doc-demo")
    assert ticket.title.startswith("What is public property")
    assert len(ticket.atoms) >= 3
    assert ticket.skills
    assert not used_llm
    assert warning == ""


def test_state_exam_profile_builds_answer_blocks() -> None:
    service = DocumentImportService()
    candidate = TicketCandidate(
        index=1,
        title="Что представляет собой государственное имущество как объект управления?",
        body=(
            "Проблема управления государственным имуществом связана с эффективным использованием публичных ресурсов. "
            "Теоретическая основа включает понятие публичного имущества, правовой режим и управленческий цикл. "
            "Практическая часть раскрывается через учет, оценку, контроль использования и выбор мер повышения эффективности. "
            "Навыки проявляются через анализ, аргументацию и подбор инструментов управления. "
            "Таким образом, имущество выступает активным управленческим ресурсом. "
            "Дополнительно можно использовать схемы и сравнение практик."
        ),
        confidence=0.9,
        section_title="state-exam",
    )
    ticket, used_llm, warning = service.build_ticket_map(
        candidate,
        "exam-demo",
        "state-exam",
        "doc-demo",
        answer_profile_code=AnswerProfileCode.STATE_EXAM_PUBLIC_ADMIN,
    )

    assert not used_llm
    assert warning == ""
    assert ticket.answer_profile_code == AnswerProfileCode.STATE_EXAM_PUBLIC_ADMIN
    assert len(ticket.answer_blocks) == 6
    assert {block.block_code for block in ticket.answer_blocks} == {
        AnswerBlockCode.INTRO,
        AnswerBlockCode.THEORY,
        AnswerBlockCode.PRACTICE,
        AnswerBlockCode.SKILLS,
        AnswerBlockCode.CONCLUSION,
        AnswerBlockCode.EXTRA,
    }
    assert not any(block.is_missing for block in ticket.answer_blocks)


def test_docx_import_smoke(tmp_path: Path) -> None:
    document_path = tmp_path / "demo.docx"
    document = Document()
    for paragraph in SOURCE_TEXT.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(document_path)

    service = DocumentImportService()
    result = service.import_document(document_path, "exam-demo", "subject-demo", "public-assets")

    assert len(result.tickets) == 2
    assert len(result.tickets[0].atoms) >= 4
    assert not result.warnings


def test_pdf_import_smoke(tmp_path: Path) -> None:
    pdf_path = tmp_path / "demo.pdf"
    pdf = canvas.Canvas(str(pdf_path))
    text = pdf.beginText(40, 800)
    for line in SOURCE_TEXT.splitlines():
        text.textLine(line)
    pdf.drawText(text)
    pdf.save()

    service = DocumentImportService()
    result = service.import_document(pdf_path, "exam-demo", "subject-demo", "public-assets")

    assert len(result.tickets) == 2
    assert len(result.tickets[1].atoms) >= 2


def test_extract_ticket_candidates_accepts_numbered_titles_without_question_mark() -> None:
    service = DocumentImportService()
    normalized = service.normalize_text(
        """1. Государственное управление в Российской Федерации. Государственное управление организует работу институтов публичной власти и задает механизм принятия решений.

2. Электронное правительство. Электронное правительство переводит услуги и коммуникацию государства в цифровую среду, повышая доступность сервисов.
"""
    )

    candidates = service.extract_ticket_candidates(normalized)

    assert len(candidates) == 2
    assert candidates[0].title == "Государственное управление в Российской Федерации"
    assert "публичной власти" in candidates[0].body
    assert candidates[1].title == "Электронное правительство"


def test_extract_ticket_candidates_uses_table_of_contents_when_present() -> None:
    service = DocumentImportService()
    normalized = service.normalize_text(
        """ОГЛАВЛЕНИЕ
1.1. Теория управления 2
1. Государственное управление в Российской Федерации. 2
2. Электронное правительство. 3

2
1.1. Теория управления
1. Государственное управление в Российской Федерации. Государственное управление организует работу институтов публичной власти, распределяет полномочия и задает механизм принятия решений.

3
2. Электронное правительство. Электронное правительство переводит услуги и коммуникацию государства в цифровую среду, повышая доступность сервисов и прозрачность процедур.
"""
    )

    candidates = service.extract_ticket_candidates(normalized)

    assert len(candidates) == 2
    assert candidates[0].section_title == "Теория управления"
    assert candidates[0].title == "Государственное управление в Российской Федерации"
    assert "институтов публичной власти" in candidates[0].body
    assert candidates[1].title == "Электронное правительство"
    assert "цифровую среду" in candidates[1].body


def test_incremental_import_preserves_saved_tickets_and_resume_finishes_tail(tmp_path: Path, monkeypatch) -> None:
    document_path = tmp_path / "demo.docx"
    document = Document()
    for paragraph in SOURCE_TEXT.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(document_path)

    facade = _build_facade(tmp_path)
    original_build = DocumentImportService.build_ticket_map
    state = {"failed_once": False}

    monkeypatch.setattr(DocumentImportService, "should_use_llm_for_structuring", lambda *args, **kwargs: False)

    def flaky_build(
        self,
        candidate,
        exam_id,
        section_id,
        source_document_id,
        ticket_id=None,
        answer_profile_code=AnswerProfileCode.STANDARD_TICKET,
    ):
        if candidate.index == 2 and not state["failed_once"]:
            state["failed_once"] = True
            raise RuntimeError("forced partial failure")
        return original_build(
            self,
            candidate,
            exam_id,
            section_id,
            source_document_id,
            ticket_id=ticket_id,
            answer_profile_code=answer_profile_code,
        )

    monkeypatch.setattr(DocumentImportService, "build_ticket_map", flaky_build)

    result = facade.import_document_with_progress(document_path)

    assert result.ok
    assert result.status == "partial_llm"
    assert result.resume_available
    assert facade.connection.execute("SELECT COUNT(*) AS total FROM tickets").fetchone()["total"] == 1
    queue_counts = facade.repository.count_import_queue_statuses(result.document_id)
    assert queue_counts["done"] == 1
    assert queue_counts["failed"] == 1

    monkeypatch.setattr(DocumentImportService, "build_ticket_map", original_build)
    resumed = facade.resume_document_import_with_progress(result.document_id)

    assert resumed.ok
    assert resumed.status == "structured"
    assert resumed.llm_done_tickets == 2
    assert resumed.resume_available is False
    assert facade.connection.execute("SELECT COUNT(*) AS total FROM tickets").fetchone()["total"] == 2
    queue_counts = facade.repository.count_import_queue_statuses(result.document_id)
    assert queue_counts["done"] == 2

    facade.connection.close()


def test_initial_import_saves_rule_based_tickets_and_queues_llm_tail(tmp_path: Path, monkeypatch) -> None:
    document_path = tmp_path / "demo.docx"
    document = Document()
    for paragraph in SOURCE_TEXT.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(document_path)

    facade = _build_facade(tmp_path, import_llm_assist=True)
    monkeypatch.setattr(DocumentImportService, "needs_llm_refinement", lambda *args, **kwargs: True)

    result = facade.import_document_with_progress(document_path)

    assert result.ok
    assert result.status == "partial_llm"
    assert result.resume_available
    assert result.tickets_created == 2
    assert facade.connection.execute("SELECT COUNT(*) AS total FROM tickets").fetchone()["total"] == 2
    queue_counts = facade.repository.count_import_queue_statuses(result.document_id)
    assert queue_counts["pending"] == 2
    assert queue_counts.get("done", 0) == 0
    facade.connection.close()


def test_finalize_import_uses_actual_ticket_count_when_document_grows(tmp_path: Path) -> None:
    document_path = tmp_path / "demo.docx"
    document = Document()
    for paragraph in SOURCE_TEXT.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(document_path)

    facade = _build_facade(tmp_path)
    result = facade.import_document_with_progress(document_path)
    source_row = facade.repository.load_source_document_row(result.document_id)

    service = DocumentImportService()
    candidate = TicketCandidate(
        index=3,
        title="Дополнительный билет",
        body="Дополнительный билет раскрывает ещё один аспект публичного управления и добавлен после базового импорта.",
        confidence=0.7,
        section_title="public-assets",
    )
    ticket, _, _ = service.build_ticket_map(candidate, "local-exam", "public-assets", result.document_id)
    facade.repository.save_ticket_map(ticket, llm_status="done", llm_error="")

    finalized = facade._finalize_import_document(result.document_id, source_row["title"], [], False)

    assert finalized.tickets_created == 3
    facade.connection.close()


def test_legacy_document_without_queue_is_marked_resumable(tmp_path: Path) -> None:
    document_path = tmp_path / "demo.docx"
    document = Document()
    for paragraph in SOURCE_TEXT.split("\n\n"):
        document.add_paragraph(paragraph)
    document.save(document_path)

    facade = _build_facade(tmp_path)
    service = DocumentImportService()
    structured = service.import_document(document_path, "local-exam", "demo", "public-assets")

    facade.repository.save_exam(
        Exam(
            exam_id="local-exam",
            title="Demo",
            description="Demo exam",
            total_tickets=len(structured.tickets),
            subject_area="demo",
        )
    )
    facade.repository.save_section(
        Section(
            section_id="public-assets",
            exam_id="local-exam",
            title="Public assets",
            order_index=1,
            description="Demo section",
        )
    )
    facade.repository.save_source_document(
        structured.source_document,
        raw_text=structured.normalized_text,
        status="structured",
        warnings=[],
        used_llm_assist=False,
        ticket_total=0,
        tickets_llm_done=0,
    )
    facade.repository.save_chunks(structured.source_document.document_id, structured.chunks)
    for ticket in structured.tickets:
        facade.repository.save_ticket_map(ticket, llm_status="done", llm_error="")

    latest = facade.load_latest_import_result()

    assert latest.status == "partial_llm"
    assert latest.resume_available
    assert latest.llm_pending_tickets == 2
    facade.connection.close()


def test_extract_ticket_candidates_uses_parenthesized_toc_entries() -> None:
    service = DocumentImportService()
    normalized = service.normalize_text(
        """ОГЛАВЛЕНИЕ
(1) 1. Государственное устройство Российской Федерации и других государств. ................................ 9
(2) 2. Выборы в федеральные органы законодательной власти. ................................ 9

9
1.1. Теоретический блок
1. Государственное устройство Российской Федерации и других государств. Государственное устройство раскрывает федеративные связи, полномочия органов и принципы публичной власти.

10
2. Выборы в федеральные органы законодательной власти. Выборы описывают порядок формирования представительных органов, стадии кампании и гарантии волеизъявления.
"""
    )

    candidates = service.extract_ticket_candidates(normalized)

    assert len(candidates) == 2
    assert candidates[0].title == "Государственное устройство Российской Федерации и других государств"
    assert "федеративные связи" in candidates[0].body
    assert candidates[1].title == "Выборы в федеральные органы законодательной власти"


def test_extract_ticket_candidates_coalesces_wrapped_toc_lines() -> None:
    service = DocumentImportService()
    normalized = service.normalize_text(
        """ОГЛАВЛЕНИЕ
(1) 1. Системный анализ ................................................................ 27
(2) 2 Основы моделирования социальных и экономических процессов. Управление
рисками ................................................................ 27
(3) 3. Контрактная система в сфере закупок товаров, работ
и услуг для государственных нужд ........................................ 52

27
1. Системный анализ. Системный анализ раскрывает структуру объекта управления, связи элементов и критерии оценки решений.

28
2. Основы моделирования социальных и экономических процессов. Управление рисками. Моделирование помогает оценивать сценарии, ограничения и последствия управленческих решений.

52
3. Контрактная система в сфере закупок товаров, работ и услуг для государственных нужд. Контрактная система описывает планирование, размещение заказа и контроль исполнения.
"""
    )

    candidates = service.extract_ticket_candidates(normalized)

    assert len(candidates) == 3
    assert candidates[1].title == "Основы моделирования социальных и экономических процессов. Управление рисками"
    assert "оценивать сценарии" in candidates[1].body
    assert candidates[2].title == "Контрактная система в сфере закупок товаров, работ и услуг для государственных нужд"


def test_recommended_import_part_count_scales_for_large_sets(tmp_path: Path) -> None:
    facade = _build_facade(tmp_path)

    assert facade._recommended_import_part_count(40) == 1
    assert facade._recommended_import_part_count(120) == 4
    assert facade._recommended_import_part_count(160) == 5
    assert facade._recommended_import_part_count(208) == 6
    facade.connection.close()


def test_import_ollama_timeout_is_unbounded_for_long_import_runs(tmp_path: Path) -> None:
    facade = _build_facade(tmp_path)

    service = DocumentImportService(
        ollama_service=facade.build_import_ollama_service(),
        llm_model=facade.settings.model,
        enable_llm_structuring=True,
    )

    assert service.ollama_service is not None
    # Теперь даже импорт использует явный generation-таймаут из настроек,
    # чтобы зависший LLM-ответ на одном билете не тормозил весь прогон.
    # Inspect-таймаут отдельный и короткий.
    assert service.ollama_service.generation_timeout_seconds == float(facade.settings.timeout_seconds)
    assert service.ollama_service.inspect_timeout_seconds == 3.0
    facade.connection.close()


def test_seed_build_can_request_unbounded_import_timeout(tmp_path: Path) -> None:
    facade = _build_facade(tmp_path)

    service = DocumentImportService(
        ollama_service=facade.build_import_ollama_service(generation_timeout_seconds=None),
        llm_model=facade.settings.model,
        enable_llm_structuring=True,
    )

    assert service.ollama_service is not None
    assert service.ollama_service.generation_timeout_seconds is None
    assert service.ollama_service.inspect_timeout_seconds == 3.0
    facade.connection.close()


def test_state_exam_import_persists_block_scores_and_statistics(tmp_path: Path) -> None:
    document_path = tmp_path / "state-exam.docx"
    document = Document()
    document.add_paragraph(
        "Билет 1. Что представляет собой государственное имущество как объект управления? "
        "Актуальность темы связана с управлением публичными ресурсами и эффективностью власти. "
        "Теоретическая часть включает понятие имущества, правовой режим и управленческий цикл. "
        "Практическая часть раскрывается через учет, оценку, контроль и выбор управленческих решений. "
        "Навыки проявляются через анализ, аргументацию и применение методов управления. "
        "В заключении имущество рассматривается как активный управленческий ресурс. "
        "Дополнительно полезны схемы и сравнительный анализ практик."
    )
    document.save(document_path)

    facade = _build_facade(tmp_path)
    result = facade.import_document_with_progress(document_path, answer_profile_code=AnswerProfileCode.STATE_EXAM_PUBLIC_ADMIN)

    assert result.ok
    assert result.answer_profile_code == AnswerProfileCode.STATE_EXAM_PUBLIC_ADMIN.value
    assert result.answer_profile_label == "Госэкзамен"
    ticket_id = facade.connection.execute("SELECT ticket_id FROM tickets LIMIT 1").fetchone()["ticket_id"]

    evaluation = facade.evaluate_answer(
        ticket_id,
        "state-exam-full",
        (
            "Проблема управления государственным имуществом связана с эффективным использованием публичных ресурсов. "
            "Теоретически важно определить правовой режим имущества и управленческий цикл. "
            "Практически нужно учитывать имущество, оценивать его использование и предлагать меры повышения эффективности. "
            "Навыки проявляются через анализ, применение методов и аргументацию решений. "
            "Вывод состоит в том, что имущество является активным ресурсом публичной власти. "
            "Дополнительно можно использовать схемы и сравнение практик."
        ),
    )

    assert evaluation.ok
    assert len(evaluation.block_scores) == 6
    assert len(evaluation.criterion_scores) == 6
    assert facade.connection.execute("SELECT COUNT(*) AS total FROM attempt_block_scores").fetchone()["total"] == 6
    assert facade.connection.execute("SELECT COUNT(*) AS total FROM ticket_block_mastery_profiles").fetchone()["total"] == 1

    statistics = facade.load_state_exam_statistics()
    assert statistics.active is True
    assert statistics.block_scores
    assert "Введение" in statistics.block_scores
    facade.connection.close()
