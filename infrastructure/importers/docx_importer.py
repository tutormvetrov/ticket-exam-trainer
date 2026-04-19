from __future__ import annotations

import logging
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from infrastructure.importers.common import ImportedDocumentText, normalize_import_title
from infrastructure.importers.ocr_support import (
    is_meaningful_image,
    ocr_image_bytes,
    should_keep_ocr_text,
)

_LOG = logging.getLogger(__name__)


def import_docx(path: str, *, workspace_root: Path | None = None) -> ImportedDocumentText:
    docx_path = Path(path)
    document = Document(str(docx_path))
    warnings: list[str] = []

    try:
        blocks = _collect_docx_blocks(document, workspace_root=workspace_root)
    except Exception as exc:
        _LOG.exception("DOCX OCR augmentation failed path=%s", docx_path)
        warnings.append(f"OCR skipped for {docx_path.name}: {exc}")
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    raw_text = "\n\n".join(block for block in blocks if block)
    return ImportedDocumentText(
        path=docx_path,
        title=normalize_import_title(docx_path.stem),
        file_type="DOCX",
        raw_text=raw_text,
        unit_count=len(document.paragraphs),
        warnings=tuple(warnings),
    )


def _collect_docx_blocks(document: Document, *, workspace_root: Path | None = None) -> list[str]:
    blocks: list[str] = []
    seen_digests: set[str] = set()

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            blocks.append(paragraph_text)

        for image_bytes in _iter_paragraph_image_bytes(document, paragraph):
            if not image_bytes or not is_meaningful_image(image_bytes):
                continue

            digest = sha256(image_bytes).hexdigest()
            if digest in seen_digests:
                continue
            seen_digests.add(digest)

            ocr_text = ocr_image_bytes(image_bytes, workspace_root=workspace_root)
            existing_text = "\n\n".join(blocks[-4:])
            if should_keep_ocr_text(existing_text, ocr_text):
                blocks.append(ocr_text)

    return blocks


def _iter_paragraph_image_bytes(document: Document, paragraph) -> list[bytes]:
    images: list[bytes] = []
    for run in getattr(paragraph, "runs", []):
        for blip in run._element.xpath('.//*[local-name()="blip"]'):
            rel_id = blip.get(qn("r:embed"))
            if not rel_id:
                continue
            image_part = document.part.related_parts.get(rel_id)
            if image_part is None:
                continue
            blob = getattr(image_part, "blob", b"")
            if blob:
                images.append(blob)
    return images
